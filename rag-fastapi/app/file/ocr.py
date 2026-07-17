"""
财报 / 通用文档 OCR 与结构化抽取模块。

统一入口：recognize(file_path, mimetype) -> str(markdown)
按 mimetype 路由到 4 条管线：
  - 图片 (image/*)            → RapidOCR 抽文字 + RapidTable 抽表格
  - PDF                        → 按页判断：电子页直接抽文，扫描页转 Pixmap 走 OCR
  - DOCX                       → 按 XML body 子节点顺序遍历段落 + 表格
  - XLSX                       → 每个 sheet 转 markdown 表格

CPU 同步任务一律 run_in_executor 包住，避免堵 FastAPI 事件循环。
OCR 引擎单例 + double-check lock 懒加载，首次请求时下载模型，不在 lifespan 预加载。
"""

import asyncio
import io
import re
import threading
from concurrent.futures import ThreadPoolExecutor
from html.parser import HTMLParser
from typing import List, Optional


DOCX_MIMETYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
XLSX_MIMETYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}
PDF_MIMETYPE = "application/pdf"

# onnxruntime session 并发不安全，单线程串行化所有 OCR 调用
_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="ocr")

_ocr_engine = None
_table_engine = None
_init_lock = threading.Lock()


def _get_ocr():
    """RapidOCR 单例，首次调用时加载模型（~15MB，下到 ~/.rapidocr/）"""
    global _ocr_engine
    if _ocr_engine is None:
        with _init_lock:
            if _ocr_engine is None:
                print("  │  [ocr] 首次加载 RapidOCR 模型（如未下载会从网络获取，约 15MB）...")
                from rapidocr_onnxruntime import RapidOCR
                _ocr_engine = RapidOCR()
                print("  │  [ocr] ✅ RapidOCR 就绪")
    return _ocr_engine


def _get_table():
    """RapidTable 单例，首次调用时加载模型（~7MB）"""
    global _table_engine
    if _table_engine is None:
        with _init_lock:
            if _table_engine is None:
                print("  │  [ocr] 首次加载 RapidTable 模型（如未下载会从网络获取，约 7MB）...")
                from rapid_table import RapidTable
                _table_engine = RapidTable()
                print("  │  [ocr] ✅ RapidTable 就绪")
    return _table_engine


def warmup():
    """供生产环境手动预热使用，默认不调"""
    _get_ocr()
    _get_table()


# ─────────────────────────────────────────────────────────────
# HTML 表格 → markdown 表格（手写，不引入 markdownify 依赖）
# ─────────────────────────────────────────────────────────────


class _HtmlTableParser(HTMLParser):
    """把 <table>...</table> 转成二维 List[List[str]]。
    不支持 colspan / rowspan 真正展开，只取单元格 text。"""

    def __init__(self):
        super().__init__()
        self.rows: List[List[str]] = []
        self._current_row: Optional[List[str]] = None
        self._current_cell: Optional[List[str]] = None

    def handle_starttag(self, tag, attrs):
        if tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._current_cell = []

    def handle_endtag(self, tag):
        if tag == "tr" and self._current_row is not None:
            self.rows.append(self._current_row)
            self._current_row = None
        elif tag in ("td", "th") and self._current_cell is not None and self._current_row is not None:
            cell_text = "".join(self._current_cell).strip()
            cell_text = re.sub(r"\s+", " ", cell_text)
            self._current_row.append(cell_text)
            self._current_cell = None

    def handle_data(self, data):
        if self._current_cell is not None:
            self._current_cell.append(data)


def _html_table_to_md(html: str) -> str:
    if not html or "<table" not in html.lower():
        return ""
    p = _HtmlTableParser()
    try:
        p.feed(html)
    except Exception:
        return ""
    rows = [r for r in p.rows if any(c for c in r)]
    if not rows:
        return ""
    # 列数对齐到最长
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("|" + "|".join(["---"] * width) + "|")
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────
# 图片管线
# ─────────────────────────────────────────────────────────────


def _boxes_to_text(ocr_result) -> str:
    """RapidOCR 输出：List[ [box, text, confidence] ]
    按 box.y 中心聚类成行，同行按 x 排序，输出多行纯文本。"""
    if not ocr_result:
        return ""
    items = []
    for entry in ocr_result:
        try:
            box, text, _ = entry
            ys = [pt[1] for pt in box]
            xs = [pt[0] for pt in box]
            y_center = sum(ys) / len(ys)
            x_start = min(xs)
            height = max(ys) - min(ys)
            items.append((y_center, x_start, height, text))
        except Exception:
            continue
    if not items:
        return ""
    items.sort(key=lambda x: (x[0], x[1]))
    avg_h = sum(it[2] for it in items) / len(items) if items else 16
    threshold = max(avg_h * 0.6, 6)
    lines: List[List] = []
    current_line: List = []
    current_y = items[0][0]
    for y, x, h, text in items:
        if not current_line:
            current_line.append((x, text))
            current_y = y
            continue
        if abs(y - current_y) <= threshold:
            current_line.append((x, text))
        else:
            current_line.sort(key=lambda t: t[0])
            lines.append([t[1] for t in current_line])
            current_line = [(x, text)]
            current_y = y
    if current_line:
        current_line.sort(key=lambda t: t[0])
        lines.append([t[1] for t in current_line])
    return "\n".join(" ".join(line) for line in lines)


def _recognize_image_array(img_input):
    """图片管线核心。img_input 可以是路径、bytes 或 numpy.ndarray，RapidOCR 都接受。
    返回 'text + markdown table' 字符串。"""
    ocr = _get_ocr()
    ocr_result, _ = ocr(img_input)
    plain = _boxes_to_text(ocr_result)

    md_table = ""
    if ocr_result:
        try:
            import numpy as np
            # rapidocr_onnxruntime 输出: List[[box(4pts), text, conf]]
            # rapid_table 期望: List[(np.ndarray boxes, Tuple[str] texts, Tuple[float] scores)]（batch）
            boxes = np.array([entry[0] for entry in ocr_result], dtype=np.float32)
            texts = tuple(str(entry[1]) for entry in ocr_result)
            scores = tuple(float(entry[2]) for entry in ocr_result)
            table_engine = _get_table()
            table_result = table_engine(img_input, ocr_results=[(boxes, texts, scores)])
            table_html = table_result.pred_htmls[0] if table_result.pred_htmls else ""
            md_table = _html_table_to_md(table_html)
        except Exception as e:
            print(f"  │  [ocr] ⚠️  表格识别失败（忽略）: {e}")
            md_table = ""

    if plain and md_table:
        return f"{plain}\n\n{md_table}"
    return plain or md_table


def _recognize_image(file_path: str) -> str:
    print(f"  │  [ocr] 路由→图片: {file_path}")
    return _recognize_image_array(file_path)


# ─────────────────────────────────────────────────────────────
# PDF 管线（电子页直接抽文字，扫描页转图走 OCR）
# ─────────────────────────────────────────────────────────────


# ─────────────────────────────────────────────────────────────
# PDF 管线（v2：财报特化）
# ─────────────────────────────────────────────────────────────
#
# 相较 v1（一律 page.get_text("text")）的增强：
#   1. 学习页眉页脚模板并从每页扣除
#   2. page.find_tables() 抽表 -> 伪表过滤 + None 填充 + 表头推断
#   3. 跨页表拼接（相邻页/列数/列宽/顶底位置四重判据）
#   4. 双栏页按 x 中心切列，避免左右串行
#   5. 电子页文本从 blocks 抽（可精确扣除表格 bbox + 页眉页脚）
#   6. 保留原有"扫描页转图 OCR"兜底路径
#
# 每一步都在 golden_set 里 5 份代表 PDF 上验证过行为合理，
# 具体阈值和判据见对应子模块的 docstring。

PDF_PAGE_TEXT_THRESHOLD = 50  # 单页文字 < 50 字视为扫描页
PDF_RENDER_DPI = 200


def _recognize_pdf(file_path: str) -> str:
    """PDF 管线 v3：财报特化 + 图像/结构增强

    Pipeline per document:
        learn_templates(doc)               # 页眉页脚模板
        for each page:
            page_text 抽取（含表 + 列 + 页眉页脚扣除）
            扫描页 或 图片主导页 -> 整页栅格化 OCR
            嵌入图 -> 独立 OCR，追加
        compute_page_sections(all_page_text)  # 目录 + 5 大分区归属
        每页前 prepend [section=xxx] marker
        merge_tables_across_pages           # 跨页表拼接
        序列化 markdown
    """
    import fitz
    from .table_extractor import extract_tables_from_page
    from .table_merger import merge_tables_across_pages
    from .column_splitter import extract_page_text_by_columns
    from .header_footer import learn_templates, normalize_hf_text
    from .image_extractor import (
        is_image_dominant_page,
        extract_embedded_images,
    )
    from .structure_detector import compute_page_sections

    print(f"  │  [ocr] 路由→PDF v3: {file_path}")
    doc = fitz.open(file_path)
    try:
        n_pages = doc.page_count
        print(f"  │  [ocr] 共 {n_pages} 页，开始学习页眉页脚模板...")
        templates = learn_templates(doc)
        print(f"  │  [ocr] 学到 {len(templates)} 个页眉页脚模板")

        page_texts: list[str] = []
        all_tables: list = []
        page_dims: dict[int, tuple[float, float]] = {}
        scan_page_ocr_count = 0
        dominant_page_ocr_count = 0
        embedded_image_ocr_count = 0

        for pno in range(n_pages):
            page = doc[pno]
            page_dims[pno] = (page.rect.width, page.rect.height)

            raw_text = (page.get_text("text") or "").strip()

            # 分支 A：扫描页（文字 <50） -> 整页 OCR
            if len(raw_text) < PDF_PAGE_TEXT_THRESHOLD:
                scan_page_ocr_count += 1
                pix = page.get_pixmap(dpi=PDF_RENDER_DPI)
                img_bytes = pix.tobytes("png")
                ocr_text = _recognize_image_array(img_bytes)
                page_texts.append(ocr_text)
                continue

            # 分支 B：图片主导页（图 ≥50% + 文本 <200） -> 整页 OCR + 保留原文本
            if is_image_dominant_page(page, len(raw_text)):
                dominant_page_ocr_count += 1
                pix = page.get_pixmap(dpi=PDF_RENDER_DPI)
                img_bytes = pix.tobytes("png")
                ocr_text = _recognize_image_array(img_bytes)
                # 图片主导页可能同时含有少量原文本；合并后再做后续处理
                page_texts.append((raw_text + "\n\n" + ocr_text).strip() if raw_text else ocr_text)
                continue

            # 分支 C：正常电子页 -> 抽表 + 抽文
            tables_on_page = extract_tables_from_page(page, pno)
            all_tables.extend(tables_on_page)
            table_bboxes = [t.bbox for t in tables_on_page if t.bbox and t.bbox != (0, 0, 0, 0)]

            body_text = extract_page_text_by_columns(page, excluded_bboxes=table_bboxes)

            if templates and body_text:
                kept_lines = [ln for ln in body_text.splitlines()
                              if normalize_hf_text(ln.strip()) not in templates]
                body_text = "\n".join(kept_lines).strip()

            # 嵌入图独立 OCR（追加到页尾，避免影响主文本布局）
            embedded_texts: list[str] = []
            try:
                for _cls, img_bytes in extract_embedded_images(page, doc):
                    embed_txt = _recognize_image_array(img_bytes)
                    if embed_txt and embed_txt.strip():
                        embedded_texts.append(embed_txt.strip())
                        embedded_image_ocr_count += 1
            except Exception as e:
                print(f"  │  [ocr] ⚠️  page {pno+1} 嵌入图 OCR 失败: {e}")

            if embedded_texts:
                body_text = (body_text + "\n\n[embedded images text]\n" +
                             "\n\n".join(embedded_texts)) if body_text else "\n\n".join(embedded_texts)

            page_texts.append(body_text)

        # 计算每页 section 归属（在跨页合并之前）
        page_infos = compute_page_sections(page_texts)
        toc_count = sum(1 for i in page_infos if i.is_toc)
        section_dist = {}
        for i in page_infos:
            k = i.section or "none"
            section_dist[k] = section_dist.get(k, 0) + 1

        # 跨页表拼接
        merged_tables = merge_tables_across_pages(all_tables, page_dims)
        n_cross = sum(1 for m in merged_tables if m.was_merged)
        print(f"  │  [ocr] 抽表 {len(all_tables)}->{len(merged_tables)} (跨页{n_cross}) "
              f"扫描页 {scan_page_ocr_count} 主导页 {dominant_page_ocr_count} "
              f"嵌入图 {embedded_image_ocr_count} 目录页 {toc_count}")
        print(f"  │  [ocr] 分区分布: {section_dist}")

        # 表格按首页归属
        table_md_by_page: dict[int, list[str]] = {}
        for m in merged_tables:
            first_page = m.source_pages[0]
            table_md_by_page.setdefault(first_page, []).append(m.to_markdown())

        # 组装最终 markdown。每页 prepend section marker（供 chunker 传到 metadata）
        parts: list[str] = []
        for pno in range(n_pages):
            info = page_infos[pno] if pno < len(page_infos) else None
            page_body = page_texts[pno] if pno < len(page_texts) else ""
            page_tables = table_md_by_page.get(pno, [])
            if not page_body and not page_tables:
                continue
            marker = ""
            if info and info.section:
                marker = f"[section={info.section} page={pno+1}]\n"
            block = marker + page_body
            if page_tables:
                block = (block + "\n\n" + "\n\n".join(page_tables)) if block else "\n\n".join(page_tables)
            parts.append(block)
        return "\n\n---\n\n".join(parts)
    finally:
        doc.close()


# ─────────────────────────────────────────────────────────────
# DOCX 管线（按 body XML 顺序遍历）
# ─────────────────────────────────────────────────────────────


def _docx_table_to_md(tbl_elem) -> str:
    from docx.oxml.ns import qn
    rows = []
    for tr in tbl_elem.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            texts = []
            for t in tc.iter(qn("w:t")):
                if t.text:
                    texts.append(t.text)
            cell_text = "".join(texts).strip()
            cell_text = re.sub(r"\s+", " ", cell_text)
            cells.append(cell_text)
        if cells:
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("|" + "|".join(["---"] * width) + "|")
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _recognize_docx(file_path: str) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    print(f"  │  [ocr] 路由→DOCX: {file_path}")
    doc = Document(file_path)
    parts: List[str] = []
    # 必须按 body 顺序遍历，doc.paragraphs / doc.tables 会丢顺序
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if text:
                parts.append(text)
        elif tag == qn("w:tbl"):
            md = _docx_table_to_md(child)
            if md:
                parts.append(md)
    return "\n\n".join(parts)


# ─────────────────────────────────────────────────────────────
# XLSX 管线
# ─────────────────────────────────────────────────────────────


XLSX_MAX_ROWS_PER_SHEET = 500
XLSX_MAX_TOTAL_CHARS = 50000


def _recognize_xlsx(file_path: str) -> str:
    from openpyxl import load_workbook
    print(f"  │  [ocr] 路由→XLSX: {file_path}")
    wb = load_workbook(file_path, data_only=True, read_only=False)
    parts: List[str] = []
    total_chars = 0
    for sheet in wb.worksheets:
        # 处理合并单元格：把合并区左上角的值填到所有覆盖位置
        merged_value = {}
        for rng in sheet.merged_cells.ranges:
            tl = sheet.cell(rng.min_row, rng.min_col).value
            for r in range(rng.min_row, rng.max_row + 1):
                for c in range(rng.min_col, rng.max_col + 1):
                    merged_value[(r, c)] = tl

        rows_data = []
        truncated = False
        for row_idx, row in enumerate(sheet.iter_rows(values_only=False), start=1):
            if row_idx > XLSX_MAX_ROWS_PER_SHEET:
                truncated = True
                break
            cells = []
            for cell in row:
                val = merged_value.get((cell.row, cell.column), cell.value)
                cells.append("" if val is None else str(val))
            rows_data.append(cells)

        if not rows_data or not any(any(c for c in r) for r in rows_data):
            continue

        width = max(len(r) for r in rows_data)
        rows_data = [r + [""] * (width - len(r)) for r in rows_data]

        sheet_lines = [f"### Sheet: {sheet.title}", ""]
        sheet_lines.append("| " + " | ".join(rows_data[0]) + " |")
        sheet_lines.append("|" + "|".join(["---"] * width) + "|")
        for r in rows_data[1:]:
            sheet_lines.append("| " + " | ".join(r) + " |")
        if truncated:
            sheet_lines.append(f"\n（此 sheet 行数过多，已截断至前 {XLSX_MAX_ROWS_PER_SHEET} 行）")
        sheet_md = "\n".join(sheet_lines)

        if total_chars + len(sheet_md) > XLSX_MAX_TOTAL_CHARS:
            parts.append("\n\n（XLSX 内容过长，后续 sheet 已截断）")
            break
        parts.append(sheet_md)
        total_chars += len(sheet_md)

    wb.close()
    return "\n\n".join(parts)


# ─────────────────────────────────────────────────────────────
# 统一入口
# ─────────────────────────────────────────────────────────────


async def recognize(file_path: str, mimetype: str) -> str:
    """统一 OCR 入口。同步任务全部丢到 _executor 串行化，不堵事件循环。"""
    loop = asyncio.get_running_loop()
    mt = (mimetype or "").lower()

    if mt.startswith("image/"):
        return await loop.run_in_executor(_executor, _recognize_image, file_path)
    if mt == PDF_MIMETYPE:
        return await loop.run_in_executor(_executor, _recognize_pdf, file_path)
    if mt in DOCX_MIMETYPES:
        return await loop.run_in_executor(_executor, _recognize_docx, file_path)
    if mt in XLSX_MIMETYPES:
        return await loop.run_in_executor(_executor, _recognize_xlsx, file_path)
    raise ValueError(f"不支持的文件类型: {mimetype}")
